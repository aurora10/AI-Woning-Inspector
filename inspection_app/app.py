import os
import re
import json
import base64
import datetime
from pathlib import Path
from io import BytesIO

from dotenv import load_dotenv
from tqdm import tqdm
from PIL import Image
from openai import OpenAI
from docx import Document
from docx.shared import Cm
from docx2pdf import convert
import imagehash

load_dotenv()

# Configuration
INPUT_DIR = Path("input_photos")
OUTPUT_DIR = Path("output")
RENAMED_DIR = OUTPUT_DIR / "renamed_photos"
PROMPT_FILE = Path("prompts/vision_prompt.txt")
DATA_FILE = OUTPUT_DIR / "inspection_data.json"
DOCX_FILE = OUTPUT_DIR / "inspection_report.docx"
PDF_FILE = OUTPUT_DIR / "inspection_report.pdf"

MAX_IMAGE_DIMENSION = 1600
SUPPORTED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}

api_key = os.environ.get("OPENAI_API_KEY")
if not api_key:
    # Try looking one level up since .env is at /Users/albert/Desktop/CodeWork/DamageAssesment
    # and we are running from inside inspection_app
    load_dotenv(dotenv_path="../.env")
    api_key = os.environ.get("OPENAI_API_KEY")

client = OpenAI(api_key=api_key)

def load_prompt():
    if not PROMPT_FILE.exists():
        raise FileNotFoundError(f"Prompt file not found: {PROMPT_FILE}")
    with open(PROMPT_FILE, "r", encoding="utf-8") as f:
        return f.read()

def get_images():
    images = []
    if not INPUT_DIR.exists():
        return images
    for path in INPUT_DIR.iterdir():
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            images.append(path)
    return sorted(images)

def sanitize_filename(name):
    name = name.lower()
    name = re.sub(r'[^a-z0-9_]', '_', name)
    name = re.sub(r'_+', '_', name)
    name = name.strip('_')
    return name[:40]

def resize_and_process_image(image_path):
    try:
        with Image.open(image_path) as img:
            # Strip EXIF by copying data
            data = list(img.getdata())
            image_without_exif = Image.new(img.mode, img.size)
            image_without_exif.putdata(data)
            
            # Resize
            img_format = img.format if img.format else "JPEG"
            if img_format == "MPO":
                img_format = "JPEG"
                
            width, height = image_without_exif.size
            if max(width, height) > MAX_IMAGE_DIMENSION:
                ratio = MAX_IMAGE_DIMENSION / max(width, height)
                new_size = (int(width * ratio), int(height * ratio))
                image_without_exif = image_without_exif.resize(new_size, Image.Resampling.LANCZOS)
                
            return image_without_exif, img_format
    except Exception as e:
        print(f"Error processing image {image_path}: {e}")
        return None, None

def image_to_base64(img, img_format):
    buffered = BytesIO()
    # Save as JPEG for consistency if format is None
    if not img_format: img_format = "JPEG"
    # Some modes cannot be saved as JPEG (like RGBA). Convert if necessary.
    if img_format.upper() == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    img.save(buffered, format=img_format)
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

def analyze_image(base64_image, prompt_text):
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt_text},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            response_format={ "type": "json_object" },
            temperature=0.1
        )
        content = response.choices[0].message.content
        return json.loads(content)
    except Exception as e:
        print(f"Vision API error: {e}")
        return None

def process_images():
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    RENAMED_DIR.mkdir(parents=True, exist_ok=True)
    
    images = get_images()
    if not images:
        print("No images found in input_photos directory.")
        return []
        
    prompt_text = load_prompt()
    results = []
    seen_hashes = set()
    
    for idx, img_path in enumerate(tqdm(images, desc="Processing images")):
        print(f"\nProcessing {img_path.name}...")
        
        # Check duplicate
        try:
            with Image.open(img_path) as img:
                img_hash = imagehash.average_hash(img)
                if img_hash in seen_hashes:
                    print(f"Skipping duplicate image: {img_path.name}")
                    continue
                seen_hashes.add(img_hash)
        except Exception as e:
            print(f"Hash calculation error for {img_path.name}: {e}")
        
        processed_img, img_format = resize_and_process_image(img_path)
        if not processed_img:
            continue
            
        base64_img = image_to_base64(processed_img, img_format)
        
        # Analyze
        analysis = analyze_image(base64_img, prompt_text)
        if not analysis:
            print(f"Failed to analyze {img_path.name}. Skipping.")
            continue
            
        room = sanitize_filename(analysis.get("room", "overige"))
        title = sanitize_filename(analysis.get("title", "onbekend"))
        
        # Determine new filename
        base_name = f"{room}_{title}"
        counter = 1
        new_filename = f"{base_name}_{counter:03d}{img_path.suffix.lower()}"
        
        while (RENAMED_DIR / new_filename).exists():
            counter += 1
            new_filename = f"{base_name}_{counter:03d}{img_path.suffix.lower()}"
            
        # Save renamed file
        if img_format.upper() == "JPEG" and processed_img.mode in ("RGBA", "P"):
             processed_img = processed_img.convert("RGB")
        processed_img.save(RENAMED_DIR / new_filename, format=img_format)
        
        print(f"[OK] Renamed to {new_filename}")
        
        record = {
            "original_filename": img_path.name,
            "new_filename": new_filename,
            "room": analysis.get("room", "Overige"),
            "description": analysis.get("description", ""),
            "category": analysis.get("category", "onduidelijk"),
            "severity": analysis.get("severity", "onbekend"),
            "confidence": analysis.get("confidence", 0.0),
            "timestamp": datetime.datetime.now().isoformat(),
            "image_width": processed_img.width,
            "image_height": processed_img.height
        }
        results.append(record)
        
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
        
    return results

def generate_docx(results):
    if not results:
        return
        
    doc = Document()
    doc.add_heading('PLAATSBESCHRIJVING – VISUELE VASTSTELLING', 0)
    doc.add_paragraph(f"Datum: {datetime.datetime.now().strftime('%d/%m/%Y')}")
    
    # Group by room
    rooms = {}
    for r in results:
        room_name = r["room"].capitalize()
        if room_name not in rooms:
            rooms[room_name] = []
        rooms[room_name].append(r)
        
    # Room order
    preferred_order = ["Inkomhal", "Woonkamer", "Keuken", "Badkamer", "Slaapkamer", "Terras", "Garage"]
    sorted_rooms = []
    
    for pref in preferred_order:
        if pref in rooms:
            sorted_rooms.append((pref, rooms.pop(pref)))
            
    # Add remaining
    overige = rooms.pop("Overige", [])
    for room, items in rooms.items():
        sorted_rooms.append((room, items))
        
    if overige:
        sorted_rooms.append(("Overige", overige))
        
    for room, items in sorted_rooms:
        doc.add_page_break()
        doc.add_heading(room, level=1)
        
        for idx, item in enumerate(items, 1):
            doc.add_heading(f"Foto {idx}", level=2)
            
            img_path = RENAMED_DIR / item["new_filename"]
            if img_path.exists():
                doc.add_picture(str(img_path), width=Cm(12))
            
            doc.add_paragraph(f"Beschrijving:\n{item['description']}")
            doc.add_paragraph(f"Categorie:\n{item['category'].capitalize()}")
            doc.add_paragraph(f"Ernst:\n{item['severity'].capitalize()}")
            doc.add_paragraph("")
            
    # Summary
    doc.add_page_break()
    doc.add_heading('SAMENVATTING', level=1)
    doc.add_paragraph(f"Totaal aantal foto's: {len(results)}")
    
    categories = {}
    for r in results:
        cat = r["category"].lower()
        categories[cat] = categories.get(cat, 0) + 1
        
    for cat, count in categories.items():
        doc.add_paragraph(f"{cat.capitalize()}: {count}")

    doc.save(DOCX_FILE)
    print(f"Generated DOCX report: {DOCX_FILE}")

def convert_to_pdf():
    if not DOCX_FILE.exists():
        print("DOCX file not found. Cannot convert to PDF.")
        return
        
    print("Converting DOCX to PDF...")
    try:
        convert(str(DOCX_FILE), str(PDF_FILE))
        print(f"Generated PDF report: {PDF_FILE}")
    except Exception as e:
        print(f"PDF Conversion failed: {e}")

if __name__ == "__main__":
    print("Starting inspection processing...")
    data = process_images()
    if data:
        generate_docx(data)
        convert_to_pdf()
        print("Done!")
    else:
        print("No images processed.")
